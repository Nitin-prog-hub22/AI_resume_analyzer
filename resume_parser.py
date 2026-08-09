"""
resume_parser.py
----------------

Responsible for reading raw text out of an uploaded resume file.
Supports .pdf and .docx formats as required by the project spec.
"""

from pathlib import Path
from io import BytesIO

from pypdf import PdfReader
from docx import Document

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE_MB = 5


class ResumeParseError(Exception):
    """Raised when a resume file cannot be read or is invalid."""


def validate_file(filename: str, file_size_bytes: int) -> None:
    """
    Validate file type and size before attempting extraction.

    Args:
        filename: original filename, used to check extension.
        file_size_bytes: size of the uploaded file in bytes.

    Raises:
        ResumeParseError: if the file type is unsupported or the file
        is too large.
    """
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeParseError(
            f"Unsupported file type '{ext}'. Please upload a PDF or DOCX file."
        )

    size_mb = file_size_bytes / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ResumeParseError(
            f"File too large ({size_mb:.1f} MB). Max allowed is {MAX_FILE_SIZE_MB} MB."
        )


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from every page of a PDF resume."""
    reader = PdfReader(BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    return "\n".join(pages_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from every paragraph (and table cell) of a DOCX resume."""
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]

    # Also pull text out of any tables, since resumes sometimes use them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """
    Detect file type from filename and dispatch to the right extractor.

    Args:
        filename: original filename (used only for extension detection).
        file_bytes: raw bytes of the uploaded file.

    Returns:
        Extracted raw text of the resume.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_text_from_docx(file_bytes)
    else:
        raise ResumeParseError(f"Unsupported file type '{ext}'.")

    if not text.strip():
        raise ResumeParseError(
            "No readable text could be extracted. The file may be a scanned "
            "image-only PDF."
        )

    return text


if __name__ == "__main__":
    # Simple manual test hook: run `python resume_parser.py sample.pdf`
    import sys

    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
        data = path.read_bytes()
        print(extract_resume_text(path.name, data)[:500])
